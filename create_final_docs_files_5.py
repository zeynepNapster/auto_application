import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re

def create_cover_letter(title, contact_info, recipient_info, letter_body, signature):
    # Create a new Document
    doc = Document()

    # Title section
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(title['name'] + "\n")
    title_run.bold = True
    title_run.font.size = Pt(16)

    title_run = title_para.add_run(title['position'])
    title_run.font.size = Pt(12)
    title_run.font.color.rgb = RGBColor(76, 175, 80)  # Green color

    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Contact Info section
    contact_info_para = doc.add_paragraph()
    for info in contact_info:
        contact_info_para.add_run(info + "\n")
    contact_info_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


    # Add a line break
    doc.add_paragraph()

    # Letter Body section
    for para in letter_body:
        doc.add_paragraph(para)

    # Signature section
    doc.add_paragraph("\nSincerely,")
    doc.add_paragraph(signature['name'])



    return doc
# Example input data
def extract_main_content(text):
    # Remove the subject line
    text_without_subject = re.sub(r"^Subject: .*\n", "", text, flags=re.MULTILINE)

    # Remove the sign-off section
    main_content = re.sub(r"\n\nBest regards,.*", "", text_without_subject, flags=re.DOTALL)

    return main_content.strip()

import os

# Import necessary modules
import os



# Example usage
file_path = 'example.txt'
replacements = {
    '(m/w/d)': '',
    '[Your Name]': '',
    'zeynep.tozge@gmail.com': 'tozgezeynep@gmail.com',
    'm/f/d':'',
    '(m/w/x)':'',
    '(f/m/d)':'',
    'm/w/d':'',
    '()':''
}
for k,folder in  enumerate(os.listdir('./applications')):
    for files in os.listdir(os.path.join('./applications',folder)):
        pth=os.path.join('./applications',folder,files)
        if 'usage' in pth:
            f = open(pth, 'r')
            desc= f.read().split('||')
            title=desc[0]
            company=desc[1]
            description=desc[2]
            url=desc[3]

        else:
            f = open(pth, 'r')
            text_data = extract_main_content(f.read())


    title = {
    "name": "zeynep tözge",
    "position": title
}

    contact_info = [
        'tozgezeynep@gmail.com',
        "+491787797311",
        "kerschlacher str 9, 81477, Munich",
        "https://www.linkedin.com/in/zeynep-t%C3%B6zge-612b5a106/"
    ]

    recipient_info = [
        "Hiring Manager"
    ]

    letter_body = [text_data]

    signature = {
        "name": "zeynep tözge"
    }
# Create the cover letter document
    doc=create_cover_letter(title, contact_info, recipient_info, letter_body, signature)
    if not os.path.exists(os.path.join(f'./final_docs/',f'application_{k}')):
        os.makedirs(os.path.join(f'./final_docs/',f'application_{k}'))

    doc.save(f'./final_docs/application_{k}/application_{k}.docx')
    f=open(f'./final_docs/application_{k}/link_{k}.txt','w')
    f.write(url)








