import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def create_cover_letter(title, contact_info, recipient_info, letter_body, signature):
    # Create a new Document
    doc = Document()

    # Title section
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(title['name'] + "\n")
    title_run.bold = True
    title_run.font.size = Pt(16)

    title_run = title_para.add_run(title['position'])
    title_run.font.size = Pt(12)
    title_run.font.color.rgb = RGBColor(76, 175, 80)  # Green color

    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Contact Info section
    contact_info_para = doc.add_paragraph()
    for info in contact_info:
        contact_info_para.add_run(info + "\n")
    contact_info_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add a line break
    doc.add_paragraph()

    # Recipient Info section
    recipient_info_para = doc.add_paragraph()
    for info in recipient_info:
        recipient_info_para.add_run(info + "\n")
    recipient_info_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Add a line break
    doc.add_paragraph()

    # Letter Body section
    for para in letter_body:
        doc.add_paragraph(para)

    # Signature section
    doc.add_paragraph("\nSincerely,")
    doc.add_paragraph(signature['name'])

    # Save the document
    file_path = "cover_letter.docx"
    doc.save(file_path)

    return file_path
# Example input data
title = {
    "name": "Jenna Beesly",
    "position": "Front Desk Receptionist"
}

contact_info = [
    "jenna@novoresume.com",
    "123 444 5555",
    "Scranton, PA",
    "linkedin.com/in/Jenna.Beesly"
]

recipient_info = [
    "Hiring Manager",
    "Permanente Luxury Inn",
    "189 Sunflower Boulevard",
    "Santa Clara, CA"
]

letter_body = [
    "Dear Front Desk Team,",
    "I am thrilled to express my interest in the front desk receptionist role at the Permanente Luxury Inn in Santa Clara. With a background in customer service and a proven track record in cash handling, I am confident in my ability to make substantial contributions to your team.",
    "Throughout my career, I have consistently prioritized creating a welcoming environment for customers. I possess a wealth of experience in verifying demographic information, streamlining check-in and check-out processes, and managing cash transactions efficiently. My adaptability to diverse situations, coupled with foundational computer skills, align seamlessly with the requirements outlined in your job listing.",
    "In my previous roles, I successfully managed multiple responsibilities, from overseeing electronic in-baskets to handling phone inquiries. This experience has honed my problem-solving and conflict resolution skills, equipping me to excel in the dynamic setting of a hotel front desk. Moreover, my familiarity with various keycard and POS systems positions me to integrate seamlessly into your existing operations.",
    "What sets me apart is my commitment to measurable results. In my previous position, I implemented a streamlined check-in process that reduced average wait times by 20%, resulting in enhanced guest satisfaction scores. Additionally, my initiatives in optimizing the use of electronic in-baskets led to a 15% increase in overall efficiency.",
    "I am confident that my skills, dedication to customer service, and proven track record of achieving tangible results can significantly contribute to the success of Permanente Luxury Inn. I am enthusiastic about the opportunity to discuss how I can further enhance your team's performance and align with your vision for customer service excellence.",
    "I welcome the chance to meet with you for an interview at your earliest convenience. Please feel free to contact me at the provided phone number or email to schedule a meeting.",
    "Thank you for considering my application. I look forward to the possibility of contributing to the continued success of Permanente Luxury Inn."
]

signature = {
    "name": "Jenna Beesly"
}

def extract_feats(doc):
    pass


for appl in os.listdir('./final_docs'):
    pth=os.path.join('./final_docs',appl)
    f=open(pth,'r')
    text_data=f.read()

# Create the cover letter document
create_cover_letter(title, contact_info, recipient_info, letter_body, signature)







